# -*- coding: utf-8 -*-
"""Generate a high-quality Word document: LOOP Matrix Developer Portal screen/route plan."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- color palette ----------
NAVY = RGBColor(0x14, 0x2A, 0x4A)
ORANGE = RGBColor(0xE8, 0x6A, 0x17)
GREY = RGBColor(0x55, 0x5B, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
AMBER = RGBColor(0xB8, 0x6A, 0x00)
RED = RGBColor(0xB0, 0x2A, 0x37)
HEADER_FILL = "142A4A"
ZEBRA_FILL = "F2F4F7"

doc = Document()

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(0x22, 0x26, 0x2B)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = WHITE
    elif color is not None:
        run.font.color.rgb = color


def style_header_row(row):
    for cell in row.cells:
        shade(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.bold = True


def add_table(headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=9.5, white=True)
    style_header_row(table.rows[0])
    status_color = {"Exists": GREEN, "Partial": AMBER, "Create": RED}
    for ridx, r in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(r):
            color = None
            bold = False
            if status_col is not None and i == status_col:
                key = val.split()[0]
                color = status_color.get(key)
                bold = True
            set_cell_text(cells[i], val, bold=bold, color=color, size=9.5)
        if ridx % 2 == 1:
            for c in cells:
                shade(c, ZEBRA_FILL)
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)
    return table


def heading(text, size=15, color=NAVY, space_before=14, space_after=6, rule=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if rule:
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "E86A17")
        pbdr.append(bottom)
        pPr.append(pbdr)
    return p


def body(text, italic=False, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p


# ================= COVER =================
for _ in range(3):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("LOOP Matrix Developer Portal")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = NAVY

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Screen Inventory, Route Map & Implementation Plan")
r.font.size = Pt(15)
r.font.color.rgb = ORANGE

s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Built on WSO2 API Manager 4.7.0 — Developer Portal (React SPA)")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY

for _ in range(8):
    doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, col in [("Document type: Technical Planning / Feasibility", GREY),
                  ("Date: 17 June 2026", GREY),
                  ("Status: Planning — no code changes", GREY)]:
    rr = meta.add_run(line + "\n")
    rr.font.size = Pt(10.5)
    rr.font.color.rgb = col

doc.add_page_break()

# ================= 1. OVERVIEW =================
heading("1.  Overview", size=16, rule=True)
body("This document catalogues every screen in the LOOP Matrix Developer Portal Figma "
     "design, maps each screen to a proposed client-side route, and classifies it by "
     "implementation status against the existing WSO2 API Manager 4.7.0 Developer Portal "
     "code base. It is a planning artifact: no application code has been modified.")
body("The portal is a client-side React 18 single-page application (SPA). New screens are "
     "added with a consistent three-step pattern, which sets the effort baseline used "
     "throughout this document:")
bullet(" — create the screen as a functional component under source/src/app/components/.", "Component")
bullet(" — register a lazy-loaded <Route> in AppRouts.jsx.", "React route")
bullet(" — add the path to WEB-INF/web.xml so direct links and refreshes are forwarded to the SPA.", "Server mapping")

heading("Status legend", size=12, space_before=10)
legend = add_table(
    ["Status", "Meaning"],
    [["Exists", "Screen already present in the code base — needs re-skinning to the LOOP design only."],
     ["Partial", "Related functionality exists and can be extended rather than built from scratch."],
     ["Create", "New screen — must be built."]],
    [1.3, 5.4], status_col=0)

# ================= 2. SUMMARY =================
heading("2.  Summary at a Glance", size=16, rule=True, space_before=16)
add_table(
    ["Metric", "Count"],
    [["Total screens", "25 (+2 shared components)"],
     ["Already exist (re-skin only)", "5  — Home, My Apps, Create App, App Detail, + 2 components"],
     ["Partially exist (extend)", "3  — Get Started docs, API Keys, App-detail tab framework"],
     ["Build new", "17 — marketing sub-pages, FAQs, full docs system, 4 app sub-tabs"]],
    [3.6, 3.1])
body("")
bullet("the Documentation system (one shared layout + content-driven pages).", "Largest single effort: ")
bullet("the product marketing landing pages (static React + MUI).", "Medium effort: ")
bullet("re-skinning the existing Home and Applications screens to the dark LOOP theme.", "Smallest effort: ")

# ================= 3. SECTION A =================
heading("3.  Screen Inventory & Proposed Routes", size=16, rule=True, space_before=16)

heading("A.  Marketing / Public Pages", size=13, color=ORANGE, space_before=10)
add_table(
    ["#", "Screen (Figma)", "Proposed route", "Status"],
    [["1", "Home — “Build Financial Infrastructure”", "/home", "Exists  (re-skin)"],
     ["2", "Developer & Resources", "/developers", "Create"],
     ["3", "Payments Landing", "/products/payments", "Create"],
     ["4", "E-commerce Landing", "/products/ecommerce", "Create"],
     ["5", "Credit Landing", "/products/credit", "Create"],
     ["6", "FAQs", "/faqs", "Create"]],
    [0.4, 2.7, 2.2, 1.4], status_col=3)

heading("B.  Documentation System (shared docs layout)", size=13, color=ORANGE, space_before=12)
add_table(
    ["#", "Doc screen", "Proposed route", "Status"],
    [["7", "API Overview", "/docs/loop-api/overview", "Create"],
     ["8", "Authorization", "/docs/loop-api/authorization", "Create"],
     ["9", "Send Money → Pesalink", "/docs/loop-api/send-money/pesalink", "Create"],
     ["10", "Send Money → M-PESA", "/docs/loop-api/send-money/mpesa", "Create"],
     ["11", "Send Money → LOOP", "/docs/loop-api/send-money/loop", "Create"],
     ["12", "Pay → M-PESA Till (Buy Goods)", "/docs/loop-api/pay/buy-goods", "Create"],
     ["13", "Pay → M-PESA Paybill", "/docs/loop-api/pay/paybill", "Create"],
     ["14", "Pay → LOOP Till", "/docs/loop-api/pay/loop-till", "Create"],
     ["15", "E-Commerce → Introduction", "/docs/loop-api/ecommerce/introduction", "Create"],
     ["16", "E-Commerce → API Doc (Create Payment Intent)", "/docs/loop-api/ecommerce/api", "Create"],
     ["17", "Get Started / Introduction", "/docs/get-started", "Partial"]],
    [0.4, 2.9, 2.6, 0.95], status_col=3)
body("Recommendation: implement items 7–17 with a single /docs/:section/:page route backed "
     "by content (Markdown / MDX) files, rather than one hard-coded route per page. This "
     "collapses 11 screens into one layout component plus content.", italic=True, size=9.5, color=GREY)

heading("C.  App / Developer Dashboard (authenticated)", size=13, color=ORANGE, space_before=12)
add_table(
    ["#", "Screen", "Proposed route", "Status"],
    [["18", "My Apps", "/applications", "Exists  (re-skin)"],
     ["19", "Create App", "/applications/create", "Exists"],
     ["20", "App Detail / Manage", "/applications/:id", "Exists"],
     ["21", "API Keys", "/applications/:id/api-keys", "Partial"],
     ["22", "Security", "/applications/:id/security", "Create"],
     ["23", "Alerts", "/applications/:id/alerts", "Create"],
     ["24", "Test Credentials", "/applications/:id/test-credentials", "Create"],
     ["25", "Go Live", "/applications/:id/go-live", "Create"]],
    [0.4, 2.3, 2.9, 1.2], status_col=3)

heading("D.  Shared Components (not routes)", size=13, color=ORANGE, space_before=12)
add_table(
    ["Component", "Status", "Notes"],
    [["User dropdown (Profile / Settings / Sign Out)", "Exists", "Header menu — re-style only."],
     ["“Ask AI” / LOOP GPT widget", "Exists", "Maps to existing Chat / ApiChat assistant."]],
    [2.9, 1.0, 2.8], status_col=1)

# ================= 4. EXISTING ROUTES =================
heading("4.  Existing Routes in the Code Base", size=16, rule=True, space_before=16)
body("For reference, the Developer Portal currently defines these client-side routes "
     "(source/src/app/AppRouts.jsx):")
for rt in ["/home", "/api-groups", "/apis  &  /api-products", "/mcp-servers",
           "/applications  (+ /create, /:id/edit, /:id, /:id/webhooks)",
           "/settings/change-password", "/search"]:
    bullet(rt)

# ================= 5. FEASIBILITY =================
heading("5.  Feasibility & Effort Assessment", size=16, rule=True, space_before=16)
body("All 25 screens are buildable within this code base using the established component / "
     "route / web.xml pattern. Effort separates into three tiers:")
add_table(
    ["Tier", "Scope", "Screens"],
    [["Low", "Re-theme existing screens to the LOOP dark look", "Home, My Apps, Create App, App Detail, shared components"],
     ["Medium", "New, largely static React + MUI pages", "Payments, E-commerce, Credit, Developer & Resources, FAQs"],
     ["High", "New mini docs engine (sidebar nav, on-this-page anchors, copyable code blocks, syntax highlighting)", "All /docs pages (items 7–17)"]],
    [0.9, 3.2, 2.6])

heading("Key considerations", size=12, space_before=10)
bullet("the LOOP Matrix visual language differs significantly from the WSO2 default; plan a dedicated theming pass.",
       "Re-skin vs. rebuild: ")
bullet("layouts port cleanly, but each doc/app screen must be wired to the correct backend data source — a decision to make per screen during build.",
       "Data wiring: ")
bullet("product names (LOOP Matrix, M-PESA, Pesalink, Paybill) are domain-specific; keep copy in i18n locale files for maintainability.",
       "Branding & content: ")

# ================= 6. NEXT STEPS =================
heading("6.  Recommended Next Steps", size=16, rule=True, space_before=16)
for i, step in enumerate([
        "Confirm the route map and naming conventions in this document.",
        "Establish the LOOP dark theme (colors, typography, spacing) as a reusable MUI theme.",
        "Build the Documentation system first — it is the highest-effort, highest-reuse area.",
        "Re-skin the existing Home and Applications screens in parallel (low risk, quick wins).",
        "Add the new marketing landing pages and FAQs.",
        "Extend the App dashboard with the Security, Alerts, Test Credentials, and Go Live sub-tabs.",
], start=1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(step)
    run.font.size = Pt(10.5)

# ---------- footer ----------
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("LOOP Matrix Developer Portal — Implementation Plan  |  Confidential  |  Page generated 2026-06-17")
fr.font.size = Pt(8)
fr.font.color.rgb = GREY

out = r"c:/2026-wso2/wso2am-4.7.0/wso2am-4.7.0/repository/deployment/server/webapps/devportal/LOOP_Matrix_DevPortal_Plan.docx"
doc.save(out)
print("SAVED:", out)
